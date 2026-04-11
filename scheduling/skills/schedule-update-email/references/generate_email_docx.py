"""
Generate a Westland Schedule Update Email as a .docx file.

Requires: pip install python-docx

Usage:
    from generate_email_docx import generate_update_email

    generate_update_email(
        output_path='Schedule Update Email - 2026-03-26.docx',
        project_info={
            'project_name': 'Project Name',
            'job_number': 'W1234',
            'contractual_completion': 'May 20, 2026',
            'projected_completion': 'December 10, 2026',
        },
        days_behind=204,           # positive = behind, negative = ahead
        gain_loss=-55,             # positive = gain, negative = loss
        successes=['Catwalks and ladders delivered and installed.'],
        gain_loss_narrative='We lost 55 days since our last update...',
        eot_recovery='Trade nonperformance has been the primary issue...',
        logic_changes='Multiple changes to logic, sequencing...',
        smartpm_changelog_url='https://live.smartpmtech.com/...',
        red_flags=['Extended durations for work that should be complete.',
                   'Rework for several trades.'],
        stalled_tasks=['Framing in each area is still not complete.',
                       'Interior HVAC and plumbing rough activities lag behind.'],
        key_items=['Material delays have been a constant concern.',
                   'Review production with OPI every single day.'],
        include_compliance_report=True,
        include_procurement_sheets=True,
        summary_screenshot_path='screenshots/smartpm-summary-report.png',
        graphs_1_screenshot_path='screenshots/smartpm-performance-graphs-1.png',
        graphs_2_screenshot_path='screenshots/smartpm-performance-graphs-2.png',
    )
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def _add_heading(doc, text, level=2):
    """Add a bold heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 2 else 11)
    return p


def _add_colored_line(doc, label, value, color):
    """Add a line with colored value text."""
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.bold = True
    run_val = p.add_run(value)
    run_val.font.color.rgb = color
    run_val.bold = True
    return p


def _add_numbered_list(doc, items, highlight_indices=None):
    """Add a numbered list with optional red highlighting."""
    highlight_indices = highlight_indices or []
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {item}')
        if i - 1 in highlight_indices:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def generate_update_email(
    output_path,
    project_info,
    days_behind=0,
    gain_loss=0,
    successes=None,
    gain_loss_narrative='',
    eot_recovery='',
    logic_changes='',
    smartpm_changelog_url='',
    red_flags=None,
    red_flag_highlights=None,
    stalled_tasks=None,
    stalled_highlights=None,
    key_items=None,
    key_item_highlights=None,
    include_compliance_report=False,
    include_procurement_sheets=False,
    summary_screenshot_path=None,
    graphs_1_screenshot_path=None,
    graphs_2_screenshot_path=None,
):
    """Generate a Westland schedule update email as a .docx file.

    Args:
        output_path: Path for the output .docx file
        project_info: Dict with keys: project_name, job_number,
                      contractual_completion, projected_completion
        days_behind: Positive = behind, negative = ahead
        gain_loss: Positive = days gained, negative = days lost
        successes: List of success strings
        gain_loss_narrative: Explanation of what drove the gain/loss
        eot_recovery: EOT / recovery efforts narrative
        logic_changes: Significant logic changes narrative
        smartpm_changelog_url: URL to SmartPM change log
        red_flags: List of red flag strings
        red_flag_highlights: List of indices (0-based) to bold/red
        stalled_tasks: List of stalled task strings
        stalled_highlights: List of indices to bold/red
        key_items: List of key items strings
        key_item_highlights: List of indices to bold/red
        include_compliance_report: Whether to mention compliance report
        include_procurement_sheets: Whether to mention procurement sheets
        summary_screenshot_path: Path to SmartPM summary report PNG (optional)
        graphs_1_screenshot_path: Path to performance graphs 1 PNG (optional)
        graphs_2_screenshot_path: Path to performance graphs 2 PNG (optional)
    """
    successes = successes or []
    red_flags = red_flags or []
    stalled_tasks = stalled_tasks or []
    key_items = key_items or []

    GREEN = RGBColor(0x00, 0x80, 0x00)
    RED = RGBColor(0xFF, 0x00, 0x00)

    doc = Document()

    # --- Project Information Header ---
    p = doc.add_paragraph()
    for label, key in [('Project', 'project_name'), ('Job Number', 'job_number'),
                       ('Contractual Completion Date', 'contractual_completion'),
                       ('Projected Substantial Completion Date', 'projected_completion')]:
        run_l = p.add_run(f'{label}: ')
        run_l.bold = True
        p.add_run(project_info.get(key, ''))
        p.add_run('\n')

    # --- Days Ahead/Behind ---
    if days_behind > 0:
        _add_colored_line(doc, 'Days Behind Schedule: ',
                          f'{days_behind} Days', RED)
    elif days_behind < 0:
        _add_colored_line(doc, 'Days Ahead of Schedule: ',
                          f'{abs(days_behind)} Days', GREEN)
    else:
        _add_colored_line(doc, 'Days Ahead/Behind Schedule: ',
                          'On Schedule', GREEN)

    # --- SmartPM Summary Report ---
    if summary_screenshot_path and os.path.isfile(summary_screenshot_path):
        doc.add_picture(summary_screenshot_path, width=Inches(6.5))
    else:
        p = doc.add_paragraph()
        p.add_run('[Insert SmartPM Summary Report screenshot here — '
                  'hyperlink to SmartPM project URL]').italic = True

    # --- Successes ---
    _add_heading(doc, 'Successes:')
    for s in successes:
        doc.add_paragraph(s, style='List Bullet')

    # --- Gain / Loss ---
    _add_heading(doc, 'Schedule Gain / Loss Since The Last Update:')
    if gain_loss > 0:
        _add_colored_line(doc, '', f'{gain_loss} Day Gain', GREEN)
    elif gain_loss < 0:
        _add_colored_line(doc, '', f'{abs(gain_loss)} Day Loss', RED)
    else:
        doc.add_paragraph('No change since last update.')

    if gain_loss_narrative:
        doc.add_paragraph(gain_loss_narrative)

    # --- EOT / Recovery ---
    _add_heading(doc, 'Status Of EOT / Recovery Efforts:')
    if eot_recovery:
        doc.add_paragraph(eot_recovery)

    # --- Significant Logic Changes ---
    _add_heading(doc, 'Significant Changes To Schedule Logic:')
    if logic_changes:
        doc.add_paragraph(logic_changes)
    if smartpm_changelog_url:
        p = doc.add_paragraph(
            'Please refer to the attached Analytics Report, '
            'or review schedule changes in SmartPM for specifics.')
        doc.add_paragraph(smartpm_changelog_url)

    # --- Red Flags ---
    _add_heading(doc, 'Red Flags:')
    _add_numbered_list(doc, red_flags, red_flag_highlights)

    # --- Stalled or Slipping Tasks ---
    _add_heading(doc, 'Stalled Or Slipping Tasks')
    _add_numbered_list(doc, stalled_tasks, stalled_highlights)

    # --- Key Items ---
    _add_heading(doc, 'Key Items & Issues To Focus On')
    _add_numbered_list(doc, key_items, key_item_highlights)

    # --- Performance Graphs Placeholder ---
    _add_heading(doc, 'Schedule Performance Graphs')
    doc.add_paragraph(
        'The charts below show our actual starts and finishes compared to planned, '
        'schedule compression, and monthly activity finish distribution. You can get '
        'a better view of these charts and drill down to greater detail regarding '
        'specific activities and trade performance by logging on to SmartPM and '
        'clicking the View Trends link on the right side of the screen.')
    if graphs_1_screenshot_path and os.path.isfile(graphs_1_screenshot_path):
        doc.add_picture(graphs_1_screenshot_path, width=Inches(6.5))
    if graphs_2_screenshot_path and os.path.isfile(graphs_2_screenshot_path):
        doc.add_picture(graphs_2_screenshot_path, width=Inches(6.5))
    if not (graphs_1_screenshot_path and os.path.isfile(graphs_1_screenshot_path)):
        p = doc.add_paragraph()
        p.add_run('[Insert SmartPM performance graph screenshots here — '
                  'hyperlink to View Trends URL]').italic = True

    # --- Compliance Report ---
    if include_compliance_report:
        doc.add_paragraph(
            'I have again included the Schedule Compliance Report in excel for your use. '
            'Please note: You will need to verify responsibility for the impacts. '
            'This report should be distributed to the Project Team each week and reviewed '
            'in detail during the OAC. Please include the form with the meeting minutes and '
            'add language to the minutes stating all parties reviewed the Schedule Compliance '
            'Report in detail and acknowledge doing so. If they wish to make any adjustments, '
            'or contest any information included in the report they may do so by responding '
            'to the meeting minutes within 24 hours, or as defined by the contract.')

    # --- Procurement Sheets ---
    if include_procurement_sheets:
        doc.add_paragraph(
            'I have included the procurement and progress update spreadsheets. '
            'Please use these to fill out all actual dates and confirmed durations '
            'prior to each update. This will significantly reduce the time we spend '
            'updating each week to give us more time to work on recovery planning.')

    # --- Closing ---
    doc.add_paragraph('Please let me know if you have any questions.')
    doc.add_paragraph('Thanks,')

    # Save
    doc.save(output_path)
    return output_path
